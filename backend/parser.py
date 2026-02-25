"""
Text extraction from PDF and DOCX files.
"""
from pathlib import Path


def extract_text(file_path: str) -> str:
    """Extract plain text from a PDF or DOCX file."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _extract_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return _extract_docx(file_path)
        else:
            return ""
    except Exception as e:
        print(f"❌ Text extraction failed for {file_path}: {e}")
        return ""


def _extract_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    pages = []

    for page in doc:
        text = page.get_text("text")
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        pages.append("\n".join(lines))

    doc.close()
    full_text = "\n\n".join(pages)

    # Remove excessive blank lines
    import re
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)
    return full_text.strip()


def _extract_docx(file_path: str) -> str:
    from docx import Document
    import re

    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)
    return full_text.strip()
