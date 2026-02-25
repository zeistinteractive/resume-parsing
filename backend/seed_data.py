"""
seed_data.py — Create and upload 10 sample resumes for testing.

Creates real DOCX files (parsed correctly by the backend), uploads them,
then waits and shows a live status table until all are done.

Usage:
    # Inside Docker (recommended — all deps already installed):
    docker compose exec backend python seed_data.py

    # Locally (API must be running on port 8000):
    pip install python-docx
    python backend/seed_data.py
"""

import io
import json
import sys
import time
import uuid
import urllib.request
import urllib.error

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌  python-docx not found.\n    Run: pip install python-docx")
    sys.exit(1)

API_BASE = "http://localhost:8000/api"

# ── 10 diverse resume profiles ────────────────────────────────────────────────

PROFILES = [
    {
        "name": "Sarah Chen",
        "email": "sarah.chen@email.com",
        "phone": "(415) 555-0101",
        "location": "San Francisco, CA",
        "title": "Senior Software Engineer",
        "summary": (
            "Senior software engineer with 7 years of experience building scalable "
            "distributed systems. Passionate about Python, cloud architecture, and "
            "machine learning infrastructure."
        ),
        "skills": [
            "Python", "FastAPI", "Django", "JavaScript", "TypeScript", "React",
            "AWS", "GCP", "Docker", "Kubernetes", "PostgreSQL", "Redis",
            "Kafka", "Terraform", "GraphQL", "TensorFlow",
        ],
        "experience": [
            {
                "company": "Stripe", "title": "Senior Software Engineer",
                "dates": "Jan 2021 – Present",
                "bullets": [
                    "Led backend development for payment processing pipeline handling $500B+ annually.",
                    "Designed microservices architecture reducing API latency by 40%.",
                    "Mentored 5 junior engineers and established code review standards.",
                ],
            },
            {
                "company": "Airbnb", "title": "Software Engineer",
                "dates": "Jun 2018 – Dec 2020",
                "bullets": [
                    "Built real-time search and recommendation system using Elasticsearch.",
                    "Improved search relevance by 25% with ML ranking models.",
                    "Developed data pipelines processing 10M events/day using Kafka.",
                ],
            },
        ],
        "education": "B.S. Computer Science — UC Berkeley, 2018 (GPA 3.9)",
    },
    {
        "name": "James Rodriguez",
        "email": "james.r@datalab.io",
        "phone": "(512) 555-0202",
        "location": "Austin, TX",
        "title": "Machine Learning Engineer",
        "summary": (
            "ML engineer with 6 years of experience building and deploying large-scale "
            "recommendation and NLP systems. Expertise in deep learning, MLOps, and "
            "statistical modelling."
        ),
        "skills": [
            "Python", "TensorFlow", "PyTorch", "Scikit-learn", "Keras",
            "Pandas", "NumPy", "Spark", "Databricks", "AWS SageMaker",
            "MLflow", "Airflow", "SQL", "R", "A/B Testing", "NLP",
        ],
        "experience": [
            {
                "company": "Meta", "title": "Senior Machine Learning Engineer",
                "dates": "Mar 2020 – Present",
                "bullets": [
                    "Built recommendation models serving 3 billion users daily.",
                    "Developed NLP content moderation models achieving 94% accuracy.",
                    "Reduced inference cost by 30% through model quantisation.",
                ],
            },
            {
                "company": "Capital One", "title": "Data Scientist",
                "dates": "Aug 2017 – Feb 2020",
                "bullets": [
                    "Created credit risk models reducing default rates by 18%.",
                    "Built real-time fraud detection system using gradient boosting.",
                    "Led A/B testing framework across 50+ product experiments.",
                ],
            },
        ],
        "education": "M.S. Statistics — UT Austin, 2017 | B.S. Mathematics — UT Austin, 2015",
    },
    {
        "name": "Priya Patel",
        "email": "priya.patel@dev.io",
        "phone": "(628) 555-0303",
        "location": "New York, NY",
        "title": "Senior Frontend Engineer",
        "summary": (
            "Creative frontend engineer with 5 years building beautiful, performant "
            "web applications. Expert in the React ecosystem, design systems, and "
            "web accessibility."
        ),
        "skills": [
            "JavaScript", "TypeScript", "React", "Next.js", "Vue.js",
            "HTML5", "CSS3", "Tailwind CSS", "GraphQL", "Jest", "Cypress",
            "Webpack", "Vite", "Figma", "Storybook", "Node.js", "WCAG 2.1",
        ],
        "experience": [
            {
                "company": "Figma", "title": "Senior Frontend Engineer",
                "dates": "Feb 2022 – Present",
                "bullets": [
                    "Built collaborative design features used by 10M+ designers worldwide.",
                    "Led JavaScript → TypeScript migration reducing bugs by 35%.",
                    "Created component library used across 8 product teams.",
                ],
            },
            {
                "company": "Shopify", "title": "Frontend Engineer",
                "dates": "May 2019 – Jan 2022",
                "bullets": [
                    "Developed merchant dashboard processing $500M in monthly transactions.",
                    "Improved Core Web Vitals scores by 60% through performance work.",
                    "Mentored 3 junior engineers and led weekly tech talks.",
                ],
            },
        ],
        "education": "B.S. Computer Science — NYU, 2019 (Minor: Graphic Design)",
    },
    {
        "name": "Michael Kim",
        "email": "michael.kim@cloudops.io",
        "phone": "(206) 555-0404",
        "location": "Seattle, WA",
        "title": "Senior DevOps / SRE Engineer",
        "summary": (
            "DevOps and site reliability engineer with 7 years managing cloud "
            "infrastructure at scale. Deep expertise in AWS, Kubernetes, and "
            "infrastructure-as-code. CKA certified."
        ),
        "skills": [
            "AWS", "GCP", "Azure", "Kubernetes", "Docker", "Terraform",
            "Ansible", "Helm", "ArgoCD", "Jenkins", "GitHub Actions",
            "Prometheus", "Grafana", "ELK Stack", "Datadog", "Python", "Go", "Bash",
        ],
        "experience": [
            {
                "company": "Amazon Web Services", "title": "Senior DevOps Engineer",
                "dates": "Apr 2020 – Present",
                "bullets": [
                    "Managed cloud infrastructure for 200+ microservices at 99.99% uptime.",
                    "Reduced infrastructure costs by $2M/year through right-sizing.",
                    "Built zero-downtime deployment pipelines used by 500+ engineers.",
                ],
            },
            {
                "company": "Microsoft", "title": "DevOps Engineer",
                "dates": "Jun 2017 – Mar 2020",
                "bullets": [
                    "Designed CI/CD pipelines for Azure DevOps platform.",
                    "Automated infrastructure provisioning with Terraform and Ansible.",
                    "Reduced deployment time from 4 hours to 15 minutes.",
                ],
            },
        ],
        "education": "B.S. Information Technology — University of Washington, 2017\nCKA | AWS Solutions Architect Professional",
    },
    {
        "name": "Lisa Wang",
        "email": "lisa.wang@pm.com",
        "phone": "(415) 555-0505",
        "location": "San Francisco, CA",
        "title": "Senior Product Manager",
        "summary": (
            "Product manager with 6 years driving product strategy at B2B SaaS companies. "
            "Strong background in user research, data analysis, and cross-functional "
            "leadership of engineering and design teams."
        ),
        "skills": [
            "Product Strategy", "Roadmap Planning", "User Research", "A/B Testing",
            "Data Analysis", "SQL", "Mixpanel", "Amplitude", "Figma", "Jira",
            "OKRs", "Agile/Scrum", "Market Research", "Go-to-Market Strategy",
        ],
        "experience": [
            {
                "company": "Salesforce", "title": "Senior Product Manager",
                "dates": "Jan 2021 – Present",
                "bullets": [
                    "Owned CRM analytics product generating $80M ARR.",
                    "Led 15 engineers and 3 designers to ship 4 major features per quarter.",
                    "Increased user engagement by 45% through data-driven improvements.",
                ],
            },
            {
                "company": "HubSpot", "title": "Product Manager",
                "dates": "Aug 2018 – Dec 2020",
                "bullets": [
                    "Launched marketing automation features adopted by 50,000+ customers.",
                    "Drove 30% increase in trial-to-paid conversion via onboarding redesign.",
                    "Defined OKRs across 3 product squads.",
                ],
            },
        ],
        "education": "MBA — Stanford GSB, 2017 | B.A. Economics — UCLA, 2014",
    },
    {
        "name": "David Thompson",
        "email": "david.t@backend.dev",
        "phone": "(312) 555-0606",
        "location": "Chicago, IL",
        "title": "Backend Engineer (Java / Spring Boot)",
        "summary": (
            "Backend engineer with 8 years of experience building enterprise-grade "
            "Java applications. Expert in Spring Boot, microservices, and high-throughput "
            "REST APIs serving millions of requests per day."
        ),
        "skills": [
            "Java", "Spring Boot", "Spring Cloud", "Hibernate", "Maven", "Gradle",
            "Kafka", "RabbitMQ", "PostgreSQL", "MySQL", "Oracle", "MongoDB",
            "Docker", "Kubernetes", "AWS", "Jenkins", "JUnit", "Mockito",
        ],
        "experience": [
            {
                "company": "JPMorgan Chase", "title": "Senior Backend Engineer",
                "dates": "Mar 2019 – Present",
                "bullets": [
                    "Architected payment API processing 50M transactions/day with 99.99% uptime.",
                    "Reduced transaction latency by 35% via connection pool tuning and caching.",
                    "Led migration from monolith to 12-service microservices architecture.",
                ],
            },
            {
                "company": "Accenture", "title": "Backend Engineer",
                "dates": "Jul 2016 – Feb 2019",
                "bullets": [
                    "Delivered Java EE applications for Fortune 500 insurance clients.",
                    "Implemented event-driven architecture using Kafka and Spring Cloud.",
                    "Reduced build times by 50% by migrating from Ant to Maven.",
                ],
            },
        ],
        "education": "B.S. Computer Science — University of Illinois Urbana-Champaign, 2016",
    },
    {
        "name": "Aisha Johnson",
        "email": "aisha.j@dataeng.io",
        "phone": "(470) 555-0707",
        "location": "Atlanta, GA",
        "title": "Senior Data Engineer",
        "summary": (
            "Data engineer with 5 years designing and maintaining large-scale data "
            "pipelines and warehouses. Expert in Apache Spark, dbt, Snowflake, and "
            "real-time streaming with Kafka."
        ),
        "skills": [
            "Python", "SQL", "Apache Spark", "dbt", "Snowflake", "BigQuery",
            "Kafka", "Airflow", "AWS Glue", "Redshift", "Delta Lake", "Iceberg",
            "Terraform", "Docker", "Great Expectations", "Looker", "Databricks",
        ],
        "experience": [
            {
                "company": "Twilio", "title": "Senior Data Engineer",
                "dates": "Jun 2021 – Present",
                "bullets": [
                    "Built real-time data platform ingesting 5B events/day using Kafka and Spark.",
                    "Reduced data warehouse query costs by 40% through partitioning strategy.",
                    "Implemented data quality framework catching 99.8% of anomalies before production.",
                ],
            },
            {
                "company": "Deloitte", "title": "Data Engineer",
                "dates": "Sep 2019 – May 2021",
                "bullets": [
                    "Designed ETL pipelines for healthcare analytics platform (HIPAA compliant).",
                    "Migrated on-premise data warehouse to Snowflake saving $400K/year.",
                    "Mentored 4 analysts on SQL optimisation and dbt modelling.",
                ],
            },
        ],
        "education": "M.S. Computer Science — Georgia Tech, 2019 | B.S. Mathematics — Spelman College, 2017",
    },
    {
        "name": "Carlos Mendez",
        "email": "carlos.m@fullstack.dev",
        "phone": "(305) 555-0808",
        "location": "Miami, FL",
        "title": "Full Stack Developer (Node.js / React)",
        "summary": (
            "Full stack developer with 4 years building end-to-end web products. "
            "Comfortable across the entire stack: React frontends, Node.js APIs, "
            "PostgreSQL databases, and AWS deployments."
        ),
        "skills": [
            "JavaScript", "TypeScript", "Node.js", "Express", "NestJS", "React",
            "Next.js", "PostgreSQL", "MongoDB", "Redis", "GraphQL", "REST APIs",
            "Docker", "AWS (EC2, RDS, S3, Lambda)", "Git", "Jest", "Prisma",
        ],
        "experience": [
            {
                "company": "Deel", "title": "Full Stack Engineer",
                "dates": "Jan 2022 – Present",
                "bullets": [
                    "Built global payroll features serving 15,000+ companies in 150 countries.",
                    "Developed real-time currency conversion service handling $2B/month.",
                    "Reduced page load time by 55% through SSR migration to Next.js.",
                ],
            },
            {
                "company": "Digital Agency XYZ", "title": "Web Developer",
                "dates": "Aug 2020 – Dec 2021",
                "bullets": [
                    "Delivered 20+ client web projects on time using React and Node.js.",
                    "Built e-commerce platform generating $5M in first-year sales.",
                    "Implemented CI/CD with GitHub Actions reducing release cycles to daily.",
                ],
            },
        ],
        "education": "B.S. Information Systems — Florida International University, 2020",
    },
    {
        "name": "Emma Williams",
        "email": "emma.w@cloudarch.io",
        "phone": "(503) 555-0909",
        "location": "Portland, OR",
        "title": "Cloud Solutions Architect (AWS)",
        "summary": (
            "AWS-certified solutions architect with 9 years designing cloud-native "
            "architectures for startups and enterprises. Specialises in serverless, "
            "cost optimisation, and multi-region disaster recovery."
        ),
        "skills": [
            "AWS (20+ services)", "Azure", "GCP", "Terraform", "CloudFormation",
            "Serverless/Lambda", "API Gateway", "EKS", "RDS", "DynamoDB",
            "VPC", "IAM", "WAF", "CloudFront", "Python", "Go",
            "FinOps", "Well-Architected Framework", "TOGAF",
        ],
        "experience": [
            {
                "company": "Salesforce", "title": "Principal Cloud Architect",
                "dates": "Sep 2019 – Present",
                "bullets": [
                    "Architected multi-region active-active platform serving 150,000 businesses.",
                    "Led cloud cost optimisation saving $8M/year through Reserved Instances and Spot.",
                    "Defined infrastructure standards adopted by 30 engineering teams.",
                ],
            },
            {
                "company": "Nike", "title": "Senior Cloud Engineer",
                "dates": "Apr 2015 – Aug 2019",
                "bullets": [
                    "Migrated 200 on-premise applications to AWS over 18 months.",
                    "Designed serverless e-commerce checkout handling Black Friday peaks (50K rps).",
                    "Achieved AWS Well-Architected compliance across all production workloads.",
                ],
            },
        ],
        "education": "B.S. Computer Engineering — Oregon State University, 2015\nAWS Solutions Architect Professional | AWS Security Specialty | TOGAF 9",
    },
    {
        "name": "Raj Sharma",
        "email": "raj.sharma@mobile.dev",
        "phone": "(669) 555-1010",
        "location": "San Jose, CA",
        "title": "Senior Mobile Developer (React Native / Flutter)",
        "summary": (
            "Mobile developer with 6 years building cross-platform iOS and Android "
            "applications. Expert in React Native and Flutter with deep experience in "
            "native integrations, performance tuning, and CI/CD for mobile."
        ),
        "skills": [
            "React Native", "Flutter", "Dart", "JavaScript", "TypeScript",
            "Swift", "Kotlin", "Redux", "MobX", "Firebase", "GraphQL",
            "Fastlane", "Xcode", "Android Studio", "Jest", "Detox",
            "App Store / Play Store", "In-app Purchases", "Push Notifications",
        ],
        "experience": [
            {
                "company": "Uber", "title": "Senior Mobile Engineer",
                "dates": "May 2021 – Present",
                "bullets": [
                    "Built driver app features used by 5M+ drivers across 70 countries.",
                    "Reduced app startup time by 45% through bundle splitting and lazy loading.",
                    "Led migration from native iOS/Android to React Native saving 60% dev time.",
                ],
            },
            {
                "company": "Swiggy", "title": "Mobile Developer",
                "dates": "Jul 2018 – Apr 2021",
                "bullets": [
                    "Developed Flutter food delivery app from scratch to 10M downloads.",
                    "Implemented real-time order tracking using WebSockets and maps SDK.",
                    "Set up Fastlane CI/CD pipelines enabling daily releases to both stores.",
                ],
            },
        ],
        "education": "B.Tech Computer Science — IIT Bombay, 2018",
    },
]


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(profile: dict) -> bytes:
    """Build a DOCX resume from a profile dict and return the raw bytes."""
    doc = Document()

    # Name
    name_para = doc.add_heading(profile["name"], level=1)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact = doc.add_paragraph(
        f"{profile['email']}  |  {profile['phone']}  |  {profile['location']}"
    )
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title_para = doc.add_paragraph(profile["title"])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].bold = True

    doc.add_paragraph()

    # Summary
    doc.add_heading("SUMMARY", level=2)
    doc.add_paragraph(profile["summary"])

    # Skills
    doc.add_heading("SKILLS", level=2)
    doc.add_paragraph("  •  ".join(profile["skills"]))

    # Experience
    doc.add_heading("EXPERIENCE", level=2)
    for job in profile["experience"]:
        p = doc.add_paragraph()
        p.add_run(f"{job['title']} — {job['company']}").bold = True
        p.add_run(f"   {job['dates']}")
        for bullet in job["bullets"]:
            doc.add_paragraph(f"• {bullet}", style="List Bullet")

    # Education
    doc.add_heading("EDUCATION", level=2)
    doc.add_paragraph(profile["education"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTTP helpers ──────────────────────────────────────────────────────────────

def _api_get(path: str) -> dict:
    with urllib.request.urlopen(f"{API_BASE}{path}", timeout=10) as r:
        return json.loads(r.read())


def _upload_docx(filename: str, docx_bytes: bytes) -> dict:
    boundary = uuid.uuid4().hex
    ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: {ctype}\r\n\r\n"
    ).encode() + docx_bytes + f"\r\n--{boundary}--\r\n".encode()

    req = urllib.request.Request(
        f"{API_BASE}/upload",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=30) as r:
        return json.loads(r.read())


# ── Status table ──────────────────────────────────────────────────────────────

STATUS_ICON = {"pending": "⏳", "success": "✅", "failed": "❌"}

def _print_table(rows: list[dict]):
    print(f"\n{'ID':<5} {'Status':<10} {'Name':<28} {'File'}")
    print("─" * 72)
    for r in rows:
        icon   = STATUS_ICON.get(r["status"], "?")
        name   = (r.get("name") or "—")[:27]
        fname  = r["filename"][:30]
        print(f"{r['id']:<5} {icon} {r['status']:<8} {name:<28} {fname}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Resume Engine — 10-resume bulk test")
    print("=" * 60)

    # 1. Health check
    try:
        _api_get("/health")
        print("✅  API is up\n")
    except Exception:
        print("❌  Cannot reach API at", API_BASE)
        print("    Start the stack with: docker compose up --build")
        sys.exit(1)

    # 2. Upload all 10 resumes
    uploaded = []
    print(f"📤  Uploading {len(PROFILES)} resumes...\n")

    for profile in PROFILES:
        filename = profile["name"].lower().replace(" ", "_") + ".docx"
        try:
            docx_bytes = build_docx(profile)
            result = _upload_docx(filename, docx_bytes)
            uploaded.append({"id": result["id"], "filename": filename,
                             "status": "pending", "name": None})
            print(f"  ✅  {filename}  →  ID {result['id']}")
        except Exception as e:
            print(f"  ❌  {filename}  →  {e}")

    if not uploaded:
        print("\n❌  No resumes uploaded. Exiting.")
        sys.exit(1)

    # 3. Poll until all done (or 3-minute timeout)
    ids = {r["id"] for r in uploaded}
    print(f"\n⏳  Waiting for Celery to parse {len(ids)} resumes "
          f"(timeout: 3 min)...\n")

    deadline = time.time() + 180
    while time.time() < deadline:
        try:
            # Fetch enough items to cover all uploaded IDs
            data = _api_get(f"/resumes?limit={len(ids) + 20}&offset=0")
            items = {r["id"]: r for r in data["items"]}
        except Exception:
            time.sleep(3)
            continue

        # Sync statuses
        rows = []
        for r in uploaded:
            item = items.get(r["id"])
            if item:
                r["status"] = item["parse_status"]
                r["name"]   = item.get("candidate_name")
            rows.append(r)

        # Clear previous table and reprint
        _print_table(rows)

        pending = sum(1 for r in rows if r["status"] == "pending")
        done    = sum(1 for r in rows if r["status"] in ("success", "failed"))

        if pending == 0:
            break

        print(f"\n  {done}/{len(ids)} done, {pending} still parsing…")
        time.sleep(4)

    # 4. Final summary
    success = sum(1 for r in uploaded if r["status"] == "success")
    failed  = sum(1 for r in uploaded if r["status"] == "failed")
    pending = sum(1 for r in uploaded if r["status"] == "pending")

    print("\n" + "=" * 60)
    print(f"  Results: ✅ {success} parsed   ❌ {failed} failed   ⏳ {pending} still pending")
    if success:
        print(f"\n  Open http://localhost:80 to browse and search resumes.")
    print("=" * 60)


if __name__ == "__main__":
    main()
